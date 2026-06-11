import fitz
import docx

def extract_text_from_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    text = ""

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                seen = []
                for item in row_text:
                    if item not in seen:
                        seen.append(item)
                text += " | ".join(seen) + "\n"

    return text

def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def extract_text(file_path: str, file_type: str) -> str:
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")